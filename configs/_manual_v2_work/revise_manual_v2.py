from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Row
from docx.text.paragraph import Paragraph


SOURCE = Path(r"D:\编程项目\phoenix\configs\6.30 洁净能源集团2026年度保温管物流链管理系统业务操作指南.docx")
OUTPUT = Path(r"D:\编程项目\phoenix\configs\8.20 洁净能源集团2026年度保温管、管件物流链管理系统业务操作指南（V2.0修编稿）.docx")


def text_of(paragraph):
    return "".join(run.text for run in paragraph.runs)


def first_run_properties(paragraph):
    for run in paragraph.runs:
        if run._r.rPr is not None:
            return deepcopy(run._r.rPr)
    return None


def set_paragraph_text(paragraph, text):
    rpr = first_run_properties(paragraph)
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, rpr)
    return paragraph


def apply_template_format(paragraph, template):
    """使用原稿同级模板重建段落属性和单一文字运行，避免多运行重排时首字符错位。"""
    current = text_of(paragraph)
    template_ppr = deepcopy(template._p.pPr) if template._p.pPr is not None else None
    rpr = first_run_properties(template)
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    if template_ppr is not None:
        paragraph._p.append(template_ppr)
    run = paragraph.add_run(current)
    if rpr is not None:
        run._r.insert(0, rpr)
    return paragraph


def find_exact(document, text):
    for paragraph in document.paragraphs:
        if text_of(paragraph) == text:
            return paragraph
    raise ValueError(f"未找到段落：{text}")


def find_starts(document, prefix):
    for paragraph in document.paragraphs:
        if text_of(paragraph).startswith(prefix):
            return paragraph
    raise ValueError(f"未找到起始段落：{prefix}")


def insert_before(anchor, template, text):
    new_p = deepcopy(template._p)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    anchor._p.addprevious(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    run = paragraph.add_run(text)
    rpr = first_run_properties(template)
    if rpr is not None:
        run._r.insert(0, rpr)
    return paragraph


def insert_block(anchor, templates, items):
    created = []
    for kind, text in items:
        created.append(insert_before(anchor, templates[kind], text))
    return created


def replace_exact(document, old, new):
    set_paragraph_text(find_exact(document, old), new)


def renumber_heading_range(document, start_heading, end_heading, old_no, new_no):
    paragraphs = list(document.paragraphs)
    start = next(i for i, p in enumerate(paragraphs) if text_of(p) == start_heading)
    end = next(i for i, p in enumerate(paragraphs) if text_of(p) == end_heading)
    for paragraph in paragraphs[start:end]:
        if paragraph.style.name in {"Heading 2", "Heading 3"}:
            current = text_of(paragraph)
            updated = re.sub(rf"^{re.escape(old_no)}\.", f"{new_no}.", current)
            if updated != current:
                set_paragraph_text(paragraph, updated)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def all_story_paragraphs(document):
    seen = set()
    for paragraph in document.paragraphs:
        if id(paragraph._p) not in seen:
            seen.add(id(paragraph._p))
            yield paragraph
    for table in document.tables:
        for paragraph in iter_table_paragraphs(table):
            if id(paragraph._p) not in seen:
                seen.add(id(paragraph._p))
                yield paragraph
    for section in document.sections:
        for story in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            for paragraph in story.paragraphs:
                if id(paragraph._p) not in seen:
                    seen.add(id(paragraph._p))
                    yield paragraph
            for table in story.tables:
                for paragraph in iter_table_paragraphs(table):
                    if id(paragraph._p) not in seen:
                        seen.add(id(paragraph._p))
                        yield paragraph


def replace_text_everywhere(document, replacements):
    for paragraph in all_story_paragraphs(document):
        current = text_of(paragraph)
        updated = current
        for old, new in replacements:
            updated = updated.replace(old, new)
        if updated != current:
            set_paragraph_text(paragraph, updated)


def clone_row(table, template_row_index=-1):
    new_tr = deepcopy(table.rows[template_row_index]._tr)
    table._tbl.append(new_tr)
    return _Row(new_tr, table)


def fill_row(row, values):
    for cell, value in zip(row.cells, values):
        paragraph = cell.paragraphs[0]
        set_paragraph_text(paragraph, value)
        for extra in cell.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)


def add_rows(table, rows):
    for values in rows:
        row = clone_row(table)
        fill_row(row, values)


doc = Document(SOURCE)

# 锁定原稿中的样式模板。
templates = {
    "h1": find_exact(doc, "一、系统概述"),
    "h2": find_exact(doc, "1.1 访问方式"),
    "h3": find_exact(doc, "3.2.1 适用角色"),
    "body": find_exact(doc, "本系统以网站形式部署，使用电脑、手机均可方便访问，网址为："),
    "item": find_exact(doc, "（1）统一保温管供应、到货、使用、损耗及计划数据口径。"),
}

# 封面与系统概述。
replace_exact(doc, "适用对象：管厂负责人、现场负责人、施工单位人员及库管人员", "适用对象：系统管理员、管厂负责人、现场负责人、施工单位人员、库管人员及只读观察人员")
replace_exact(doc, "版本：V1.0", "版本：V2.0")
replace_exact(doc, "编制日期：2026年7月1日", "修编日期：2026年8月20日")
insert_before(
    find_exact(doc, "1.2 系统用途"),
    templates["body"],
    "使用手机访问时，页面会根据屏幕宽度调整为卡片或紧凑列表。涉及发货、到货、接收、库管确认和现场填报等操作时，应先核对当前施工标段、业务日期和物资类别，再进行提交。",
)
replace_exact(
    doc,
    "《洁净能源集团2026年度保温管物流链管理系统》用于城市管网更新项目，实现保温管供应、发货、到货、接收、库存、消耗及滚动计划管理。",
    "《洁净能源集团2026年度保温管、管件物流链管理系统》用于城市管网更新项目，实现保温管和管件的供应、发货、到货、接收、库存、使用及计划管理。",
)
replace_exact(
    doc,
    "系统围绕保温管物流链条，建立从管厂发货、现场到货、施工接收、库管监督到每日消耗填报的全过程在线台账，实现保温管物资流转信息共享、过程留痕、数据统计和供应保障分析等功能。",
    "系统围绕保温管和管件物流链条，建立从供给侧发货、现场到货、施工接收、库管监督到每日使用填报的全过程在线台账，实现物资流转信息共享、过程留痕、数据统计和供应保障分析等功能。",
)
replace_exact(doc, "（1）统一保温管供应、到货、使用、损耗及计划数据口径。", "（1）统一保温管和管件供应、到货、接收、库存、使用及计划数据口径。")
replace_exact(doc, "（3）支撑管厂根据现场需求和三日净缺口组织生产、发运。", "（3）支撑供给侧根据现场需求、三日净缺口和管件采购计划组织生产、发运。")
replace_exact(doc, "（1）全链条协同。以“供给侧、需求侧、库管侧”共同参与为基础，将发货、到货、接收、监督、消耗和计划纳入同一业务链条。", "（1）全链条协同。以“供给侧、需求侧、库管侧”共同参与为基础，将保温管和管件的发货、到货、接收、监督、使用和计划纳入同一业务链条。")

# 角色表与基础数据。
role_table = doc.tables[0]
replace_exact(doc, "系统按照物流链管理需要，主要分为系统管理员、供给侧、需求侧和库管侧四类角色。", "系统按照物流链管理需要，主要分为系统管理员、供给侧、需求侧、库管侧和观察角色五类角色。")
fill_row(role_table.rows[1], ["系统管理员", "系统管理员", "负责系统基础配置、账号权限维护、核心日期维护、提交情况核查、异常协助处理和操作记录查询等工作"])
fill_row(role_table.rows[2], ["供给侧", "管厂负责人", "负责保温管和管件发货登记、发货记录核对、供给侧数据查看和流转凭证查询等工作"])
fill_row(role_table.rows[3], ["需求侧", "现场负责人", "负责保温管和管件到货确认、保温管每日使用量和损耗量填报、三日滚动计划维护、管件安装使用量填报等工作"])
fill_row(role_table.rows[4], ["", "施工单位人员", "负责对已到货保温管和管件进行接收确认，发现数量差异时按页面要求填写实际数量和原因说明"])
fill_row(role_table.rows[5], ["库管侧", "库管人员", "负责查看保温管和管件的到货、接收、库存和流转状态，进行批量库管确认或整车归档，发挥信息知晓和监督作用"])
add_rows(role_table, [["观察角色", "全局观察员", "按照只读方式查看项目各业务页面、历史记录和统计结果，不进行业务数据提交"]])
# 原表“需求侧”为跨两行合并单元格；逐行回填后需恢复合并单元格标题。
set_paragraph_text(role_table.cell(3, 0).paragraphs[0], "需求侧")
insert_before(find_exact(doc, "2.3 基础数据准备"), templates["body"], "页面右上角可显示当前在线情况。该信息用于了解协同人员是否在线，不替代电话、工作群等现场沟通方式。")
replace_exact(doc, "（1）管厂、换热站、管材规格列表等基础信息维护。", "（1）供给主体、施工标段、保温管规格、管件类型和计量单位等基础信息维护。")
replace_exact(doc, "（3）各换热站、各规格保温管初始设计使用量、计划使用量准备。", "（3）各施工标段保温管设计使用量、计划采购量以及管件设计量、计划采购量准备。")
replace_exact(doc, "（4）系统正式运行前连续三日计划量录入。", "（4）系统正式运行前连续三日保温管计划量录入，并核对管件基础台账。")
replace_exact(
    doc,
    "上线初期，系统管理员将开放三列计划量填报表格，用于一次性录入连续三日计划。正常运行后，系统每日仅开放最后一列计划量，由现场负责人补充新的第三日计划。",
    "系统管理员根据实际管理阶段设置计划可修改日期。页面中处于可编辑状态的日期可以填报或调整，锁定日期不可修改；现场负责人应以页面显示的可编辑范围为准，补齐连续三日计划。",
)

# 原章节重命名、顺延编号。
set_paragraph_text(find_exact(doc, "三、业务流程"), "三、保温管业务流程")
renumber_heading_range(doc, "四、数据看板口径", "五、核心指标口径与异常处理", "4", "5")
renumber_heading_range(doc, "五、核心指标口径与异常处理", "六、常见问题", "5", "6")
renumber_heading_range(doc, "六、常见问题", "七、填报纪律与管理要求", "6", "7")
set_paragraph_text(find_exact(doc, "四、数据看板口径"), "五、数据看板口径与查看")
set_paragraph_text(find_exact(doc, "五、核心指标口径与异常处理"), "六、核心指标口径、异常处理及历史查询")
set_paragraph_text(find_exact(doc, "六、常见问题"), "七、常见问题")
set_paragraph_text(find_exact(doc, "七、填报纪律与管理要求"), "九、填报纪律与管理要求")
set_paragraph_text(find_exact(doc, "八、核心操作速查表"), "十、核心操作速查表")

# 保温管业务流程优化。
replace_exact(doc, "系统主流程包括五个环节：管厂发货登记、现场负责人确认到货、施工单位人员确认接收、库管人员确认知晓、现场负责人每日填报使用量、损耗量和滚动计划。", "保温管主流程包括五个环节：管厂发货登记、现场负责人确认到货、施工单位人员确认接收、库管人员确认知晓、现场负责人每日填报使用量、损耗量和滚动计划。")
insert_before(find_exact(doc, "3.2.4 注意事项"), templates["body"], "发货前，管厂负责人可根据施工标段、管材规格、三日计划和净缺口等信息筛选待发物资。筛选结果用于辅助组织发运，最终发货数量仍应以实际装车情况为准。")
insert_block(
    find_exact(doc, "3.3 需求侧操作——现场负责人确认到货"),
    templates,
    [
        ("h3", "3.2.5 流转凭证查看"),
        ("body", "在发货记录中点击状态或“流转凭证”，可查看车次、订单、车牌号、供需主体以及发货、到货、施工接收、库管确认等节点信息。查看凭证不改变业务状态。"),
    ],
)
insert_block(
    find_exact(doc, "3.6 每日填报：使用量、损耗量与三日滚动计划"),
    templates,
    [
        ("h3", "3.5.4 查看方式与批量确认"),
        ("body", "库管页面支持按单条明细查看，也可按车次汇总查看。需要处理多条待库管记录时，可勾选符合条件的记录，统一填写批量确认备注后提交；管件业务按照整车进行归档。"),
        ("h3", "3.5.5 数据导出"),
        ("body", "库管人员可根据当前筛选结果导出 Excel 台账。导出前应先确认业务类别、施工标段、日期和状态筛选条件是否符合核查需要。"),
    ],
)
replace_exact(doc, "业务日期通常为实际填报日的前一日，即D-1。", "填报页面显示的“消耗采集日期”为本次业务日期。系统启用自动日期时，北京时间每日6:30完成业务日切换；通常情况下，消耗采集日期为当前业务日的前一日，即D-1。")
replace_exact(doc, "例如：6月30日8:00前进行填报时，业务日期为6月29日。现场负责人需要填报6月29日实际使用量和实际损耗量。", "例如：北京时间6月30日6:30后，页面显示的消耗采集日期通常为6月29日。现场负责人应以页面实际显示日期为准，填报对应日期的使用量和损耗量。")
insert_block(
    find_exact(doc, "四、管件业务流程") if any(text_of(p) == "四、管件业务流程" for p in doc.paragraphs) else find_exact(doc, "五、数据看板口径与查看"),
    templates,
    [
        ("h3", "3.6.8 Excel批量粘贴录入"),
        ("body", "三日滚动计划可从 Excel 复制“型号、计划量”数据块，在页面批量粘贴区按 Ctrl+V 录入；每日使用量和损耗量可复制“型号、使用量、损耗量、备注”数据块进行匹配填充。"),
        ("body", "批量粘贴后，应逐行核对型号匹配结果和数量。未匹配、重复或格式不正确的数据应在页面中调整后再提交。"),
        ("h3", "3.6.9 提交与页面切换"),
        ("body", "表格填写完成后，应点击相应提交按钮并确认系统提示成功。若表格中存在未保存修改，切换标签、切换施工标段或离开页面时，系统会提示是否放弃修改；选择离开后，未提交内容不会保留。"),
        ("h3", "3.6.10 可填日期与填报顺序"),
        ("body", "现场负责人只能填写页面中已开放的日期和单元格。系统启用严格计划填报流程时，应先完成前一日使用量和损耗量填报，再填写新增第三日计划；未启用时，两项业务可以分别提交。"),
    ],
)

# 新增管件业务流程，插入到数据看板章节之前。
chapter5_anchor = find_exact(doc, "五、数据看板口径与查看")
fitting_items = [
    ("h1", "四、管件业务流程"),
    ("h2", "4.1 业务流程总览"),
    ("body", "管件业务按照整车组织流转，主要包括供给侧整车发货、现场负责人确认到货、施工单位确认接收、库管人员整车归档、现场负责人填报安装使用量五个环节。"),
    ("body", "同一车辆可以装载多种管件。系统按车次汇总展示，同时保留每种管件的类型、型号规格、数量、单位和订单编号，便于各方按整车核对、按明细追溯。"),
    ("h2", "4.2 供给侧操作——管件整车发货登记"),
    ("h3", "4.2.1 适用角色"),
    ("body", "管厂负责人。"),
    ("h3", "4.2.2 操作时点"),
    ("body", "管件装车完成并准备发运时，管厂负责人应在“供给侧管理入口”的“管件业务”中完成整车发货登记。"),
    ("h3", "4.2.3 操作方式"),
    ("item", "（1）进入“管件发货与明细记录”，填写运输车牌号、供给主体、目标施工标段和整车备注。"),
    ("item", "（2）在本车管件发货表格中逐行填写管件类型、型号或规格、发货数量、单位和备注。同一车辆的不同管件应填写在同一整车发货单中。"),
    ("item", "（3）管件明细较多时，可先下载标准填报模板，在 Excel 中完成填写后再复制或导入页面。模板中的必填列、单位和数据格式不得随意调整。"),
    ("item", "（4）填写完成后，逐行核对管件类型、型号规格、数量和单位。发货数量应大于0，单位应使用系统允许的单位。"),
    ("item", "（5）如页面提示存在非常用或异形管件，应再次核对名称和规格。确认属于实际发货物资后，可以继续提交；如为录入错误，应返回修改。"),
    ("item", "（6）点击“提交整车管件发货单”。提交成功后，系统生成管件车次号和对应订单号，并在已提交管件发货记录中显示。"),
    ("h3", "4.2.4 同车牌近期车次处理"),
    ("body", "系统发现同一车牌在1小时内已有管件发货记录时，会提示是否将本次明细合并到前序车次。管厂负责人应核对前序车次号、车牌号和已装管件，再决定继续合并或返回修改车牌。"),
    ("body", "确认合并后，本次新增管件与前序记录共用同一车次号，作为同一实际车辆统一流转；不属于同一车辆的，不得为了减少车次而合并。"),
    ("h3", "4.2.5 注意事项"),
    ("item", "（1）车牌号为整车识别的重要依据，提交前必须核对。"),
    ("item", "（2）同一车辆存在多种管件时，应一次性录入或在1小时合并提示中确认追加，避免拆成多个无关车次。"),
    ("item", "（3）标准类型提示用于统一名称，不影响实际存在的非常用管件登记；非常用管件应填写清楚型号规格和备注。"),
    ("item", "（4）提交后如发现发货信息错误，应按页面允许方式处理；记录已进入后续确认流程时，应联系系统管理员协助核查。"),
    ("h2", "4.3 需求侧操作——管件到货确认"),
    ("h3", "4.3.1 适用角色"),
    ("body", "现场负责人。"),
    ("h3", "4.3.2 操作时点"),
    ("body", "管件车辆实际到达指定施工标段并完成现场清点后，现场负责人应在“需求侧管理入口”的“管件发货记录”中进行到货确认。"),
    ("h3", "4.3.3 操作方式"),
    ("item", "（1）按发货日期、车牌号、车次号、管件类型或型号查询待到货车次。"),
    ("item", "（2）展开车次，核对供给主体、车牌号、整车备注和各项管件明细。"),
    ("item", "（3）点击“到货确认”，逐项填写实际到货数量；实际到货数量不得大于对应发货数量。"),
    ("item", "（4）如实际到货数量与发货数量不一致，应按实际清点数量填写，并在到货备注中说明短缺、拆包或其他现场情况。"),
    ("item", "（5）核对无误后，点击“确认物理卸车到货”。提交成功后，车次状态进入待施工接收。"),
    ("h2", "4.4 需求侧操作——施工单位确认接收"),
    ("h3", "4.4.1 适用角色"),
    ("body", "施工单位人员。"),
    ("h3", "4.4.2 操作时点与方式"),
    ("body", "现场负责人完成到货确认后，施工单位人员应进入“管件发货记录”，找到状态为“待施工接收”的车次，核对各项管件实际到货数量和现场领用情况。"),
    ("body", "确认无误后，点击“施工接收”，填写施工领用接收说明或差异说明，再点击“确认施工无误接收”。提交成功后，车次进入待库管归档。"),
    ("h3", "4.4.3 注意事项"),
    ("body", "施工接收应以现场实际清点和领用情况为准。发现管件类型、型号规格或数量存在问题时，应先与现场负责人核对，并在接收说明中写明情况，不得在未核实的情况下直接确认。"),
    ("h2", "4.5 库管侧操作——整车归档"),
    ("h3", "4.5.1 适用角色"),
    ("body", "库管人员。"),
    ("h3", "4.5.2 操作方式"),
    ("item", "（1）进入“库管管理入口”，切换到“管件发货记录”。"),
    ("item", "（2）查看待现场到货、待施工接收、待库管确认和已入库完结等状态汇总。"),
    ("item", "（3）找到待库管归档车次，展开并核对车牌号、供需主体、发货数量、实际到货数量、施工接收状态和整车备注。"),
    ("item", "（4）确认无误后，点击“整车批量归档”。归档完成后，该车次状态显示为“库管已归档”。"),
    ("body", "库管人员可导出管件台账，也可点击“流转凭证”查看整车从发货到归档的完整记录。"),
    ("h2", "4.6 管件库存与安装使用量填报"),
    ("h3", "4.6.1 适用角色"),
    ("body", "现场负责人。"),
    ("h3", "4.6.2 填报日期"),
    ("body", "页面显示的“消耗采集日期”为本次管件安装使用量对应日期。现场负责人应先核对日期，再填写本日实际安装使用数量。"),
    ("h3", "4.6.3 操作方式"),
    ("item", "（1）进入“需求侧管理入口”的“管件业务”，选择“库存与管件使用量填报”。"),
    ("item", "（2）查看到货物料种类、累计到货总数、累计安装总数、现场实时可用库存和整体安装消耗率。"),
    ("item", "（3）按管件类型和型号规格填写本日使用量，可以直接输入，也可以使用加减按钮、滑块或“全部用完”快捷操作。"),
    ("item", "（4）本日使用量应为实际安装数量，不得大于页面显示的现场可用库存；无库存的管件不能填报使用量。"),
    ("item", "（5）需要说明施工班组、使用部位或其他情况时，可填写备注。核对合计数量后提交。"),
    ("h3", "4.6.4 提交规则与历史查看"),
    ("body", "同一施工标段在同一消耗采集日期只提交一次管件安装使用记录。页面显示“单日填报已锁定”时，不得重复提交。"),
    ("body", "已提交记录按日期归入管件现场安装使用历史台账。操作人员可查看每日明细并导出 Excel；发现数据问题时，应联系系统管理员核查处理。"),
    ("h2", "4.7 管件设计量与计划采购量查询"),
    ("body", "管厂负责人、现场负责人和管理人员可在各自业务页面查看管件设计量与计划采购量台账，并按施工标段、管件类型、型号规格等条件筛选。"),
    ("body", "台账主要用于采购准备、发货组织和现场核对。需要线下分析时，可导出 Excel。查询人员应先确认当前施工标段，避免将不同标段数据混合使用。"),
    ("h2", "4.8 流转凭证与台账导出"),
    ("body", "在供给侧、需求侧和库管侧的管件车次卡片中，点击“流转凭证”可查看车牌号、整车管件清单、发货、到货、施工接收和库管归档等节点信息。"),
    ("body", "各业务页面的导出文件均为 Excel（.xlsx）格式。导出前应先设置日期、施工标段、状态或关键字筛选条件，导出内容以当前查询范围为准。"),
]
insert_block(chapter5_anchor, templates, fitting_items)

# 数据看板日期口径，保留原章表达方式但改为用户可见规则。
insert_before(find_exact(doc, "5.2 双日期口径"), templates["body"], "本章所称库存、计划和净缺口指标主要用于保温管统计。管件库存和安装使用情况在需求侧“库存与管件使用量填报”页面查看，按管件实际单位统计。")
replace_exact(doc, "业务日期用于控制数据录入。每日上午8:00前填报业务日期实际使用量、损耗量，绑定至填报行为发生日的前一日。也就是说，实际填报日为D时，实际使用量和损耗量对应业务日期D-1。", "业务日期用于控制数据录入。操作人员应以页面显示的“消耗采集日期”为准填报实际使用量和损耗量。系统启用自动日期时，北京时间每日6:30换日；通常实际业务日为D时，消耗采集日期对应D-1。")
replace_exact(doc, "每日现场负责人完成前一日实际使用量、损耗量和新一日计划量填报后，系统管理员对数据进行核查。确认无误后，系统管理员在后台推进展示日期。", "核心日期可以由系统管理员手动维护，也可以按页面设置自动更新。现场负责人完成实际使用量、损耗量和计划量填报后，系统管理员应结合提交状态和数据完整性进行核查。")
replace_exact(doc, "展示日期推进后，看板、大盘汇总和供给侧分析的数据窗口整体向前滚动一天。新一天的到货、使用量和损耗量将在同一业务时间轴上被纳入统计。", "选择“否”时，展示日期、消耗采集日期和计划起始日期均由系统管理员维护；选择“是”时，计划起始日期和消耗采集日期自动更新，展示日期仍由系统管理员维护；选择“全部是”时，三个日期均按北京时间6:30业务日边界自动更新。")

# 核心指标表补充管件口径。
metric_table = doc.tables[1]
add_rows(
    metric_table,
    [
        ["管件发货数量", "管厂确认发货的管件数量，按系统中该管件对应单位统计。"],
        ["管件到货数量", "现场负责人清点并确认的实际到货数量，不得大于发货数量。"],
        ["管件累计安装数量", "各消耗采集日期内已提交的管件实际安装使用数量合计。"],
        ["管件现场可用库存", "累计实际到货数量减累计安装使用数量；填报时不得超过当前可用库存。"],
        ["管件安装消耗率", "累计安装数量与累计到货数量的比例，用于查看现场管件使用进度。"],
    ],
)
insert_block(
    find_exact(doc, "6.3 历史查询与数据导出"),
    templates,
    [
        ("h3", "6.2.7 管件同车牌合并提示"),
        ("body", "如页面提示同一车牌在1小时内已有管件发货记录，应先判断是否属于同一实际车辆。属于同一车辆的，可以确认合并；不属于同一车辆或车牌填写有误的，应返回修改。"),
        ("h3", "6.2.8 管件实际到货数量小于发货数量"),
        ("body", "现场负责人应按实际清点数量填写到货确认数，并在到货备注中说明差异。实际到货数量不得大于发货数量。"),
        ("h3", "6.2.9 管件使用量超过库存"),
        ("body", "如系统提示管件使用量超过库存，应核对管件类型、型号规格、当前库存和填报数量。确认有到货记录尚未完成到货确认的，应先完成到货确认；仍有问题的，应联系系统管理员核查。"),
    ],
)
replace_exact(doc, "6.3 历史查询与数据导出", "6.3 历史查询与Excel导出")
replace_exact(
    doc,
    "系统支持历史数据查询与导出。管理人员可根据日期、换热站、规格、状态等条件查询发货、到货、接收、库存、使用量、损耗量和计划数据。导出文件采用CSV格式，主要考虑兼容性和文件大小。系统导出数据主要用于过程管理、供应调度、内部核对和统计分析。",
    "历史查询页面分为“保温管历史数据”和“管件发货历史数据”两个标签。保温管可按日期、施工标段等条件查询计划量、使用量、损耗量、到货量和在途情况，并查看分规格明细与时段汇总；管件可按日期、施工标段、车牌号、管件类型、型号规格和状态等条件查询整车及明细记录。查询完成后，点击“导出台账（.xlsx）”生成 Excel 文件。导出数据主要用于过程管理、供应调度、内部核对和统计分析。",
)

# 常见问题新增用户高频场景。
chapter9_anchor = find_exact(doc, "九、填报纪律与管理要求")
faq_items = [
    ("h2", "7.13 一辆车有多种管件，是否需要分开提交？"),
    ("body", "不需要。同一车辆装载的不同管件应在同一整车发货单中逐行填写，统一提交并生成同一车次。"),
    ("h2", "7.14 为什么提交管件发货时提示1小时内有同车牌记录？"),
    ("body", "系统判断同一车牌在1小时内已有发货记录，提示核对是否属于同一实际车辆。属于同一车辆的，可以确认合并；不属于同一车辆的，应返回修改车牌或重新核对。"),
    ("h2", "7.15 管件实际到货数量比发货数量少，怎么处理？"),
    ("body", "现场负责人应在到货确认时按实际清点数量填写，并在到货备注中说明差异。系统按实际确认数量记录到货。"),
    ("h2", "7.16 为什么管件使用量不能再次提交？"),
    ("body", "同一施工标段在同一消耗采集日期只允许提交一次。页面显示“单日填报已锁定”时，说明当天记录已经提交；如数据有误，应联系系统管理员核查处理。"),
    ("h2", "7.17 手机上可以进行发货、到货和使用量填报吗？"),
    ("body", "可以。手机端会将部分表格调整为卡片或紧凑列表。操作时应先确认施工标段、业务日期和车次，再填写并提交。"),
    ("h2", "7.18 历史数据从哪里导出？"),
    ("body", "进入“历史查询”，选择保温管或管件标签，设置查询条件并完成查询后，点击“导出台账（.xlsx）”。供给侧、需求侧和库管侧部分业务台账也支持按当前筛选结果导出 Excel。"),
]
insert_block(chapter9_anchor, templates, faq_items)

# 新增系统管理员日常操作章节。
admin_items = [
    ("h1", "八、系统管理员日常操作"),
    ("h2", "8.1 核心日期维护"),
    ("h3", "8.1.1 操作入口"),
    ("body", "系统管理员进入“全局管理”的核心参数区域，查看展示日期、消耗采集日期、滚动计划起始日期、计划可修改天数和计划填报流程设置。"),
    ("h3", "8.1.2 日期更新方式"),
    ("item", "（1）选择“否”：三个日期均由系统管理员根据业务进度手动维护。"),
    ("item", "（2）选择“是”：滚动计划起始日期和消耗采集日期自动更新，展示日期仍由系统管理员维护。"),
    ("item", "（3）选择“全部是”：三个日期均自动更新。自动日期以北京时间每日6:30为业务日切换边界。"),
    ("body", "保存前应核对页面显示的三个日期。日期变化会影响现场填报日期、计划窗口和看板统计口径，不得在未确认业务进度的情况下随意调整。"),
    ("h2", "8.2 计划可修改范围与填报顺序"),
    ("body", "“计划可填报修改天数”用于确定三日计划中可编辑的日期数量。系统管理员应根据项目阶段设置，现场负责人以页面实际开放的单元格为准。"),
    ("body", "“严格计划填报流程管控”开启时，现场负责人应先完成前一日使用量和损耗量填报，再填新增第三日计划；关闭时，两项业务可以分别填报。调整后应及时通知相关现场负责人。"),
    ("h2", "8.3 提交状态核查"),
    ("body", "系统管理员可在全局管理中查看各施工标段的提交状态、最近提交日期、完成时间和填报人员。发现未提交、日期不一致或长时间未更新时，应联系对应现场负责人核查。"),
    ("h2", "8.4 基础参数与基准台账维护"),
    ("body", "系统管理员负责维护供给主体、施工标段、保温管规格、管件允许单位、常用标准管件类型以及保温管、管件设计量和计划采购量等基础信息。"),
    ("body", "新增或调整基础数据前，应先与业务负责人确认名称、规格、单位和适用施工标段。保存后，应在供给侧和需求侧页面核对是否正确显示。"),
    ("h2", "8.5 业务记录与操作记录查询"),
    ("body", "系统管理员可根据日期、账号、业务类型和操作结果查询提交记录或操作记录，用于核查业务处理过程。查询记录主要用于事实核对，不替代现场签认和业务责任确认。"),
    ("h2", "8.6 异常协助处理"),
    ("body", "接到用户反馈后，系统管理员应先核对施工标段、业务日期、车次号或订单号、物资类型、当前状态和页面提示，再决定处理方式。涉及已进入后续流程的数据，不应仅凭口头说明直接调整。"),
]
insert_block(chapter9_anchor, templates, admin_items)

# 填报纪律和速查表。
replace_exact(doc, "（5）现场负责人应在每日上午8:00前完成规定数据填报。", "（5）现场负责人应在每日上午8:00前完成保温管规定数据填报，并按页面显示的消耗采集日期及时填报管件安装使用量。")
replace_exact(doc, "（6）管厂负责人应及时、准确登记发货信息。", "（6）管厂负责人应及时、准确登记保温管和管件发货信息；同一车辆的管件应按整车组织填报。")
replace_exact(doc, "（8）库管人员应定期查看和确认相关记录，发挥监督管理作用。", "（8）库管人员应定期查看保温管和管件记录，及时完成批量库管确认或整车归档，发挥监督管理作用。")
replace_exact(doc, "（9）各类数量填报应以米为单位，按实际情况填写。", "（9）保温管数量以米为单位；管件数量按页面显示单位填写。各类数量均应以现场实际情况为准。")

quick_table = doc.tables[2]
existing_quick_rows = [
    ["保温管发货登记", "管厂负责人", "车辆驶离时", "核对施工标段、规格、数量、车牌号后确认发货"],
    ["保温管发货撤回", "管厂负责人", "确认到货前", "发货记录有误时撤回并重新发货"],
    ["保温管确认到货", "现场负责人", "车辆实际到站后第一时间", "按实际到货型号、数量确认，确认后进入库存计算"],
    ["保温管接收确认", "施工单位人员", "确认到货后尽快", "12小时未操作的，系统自动按确认到货量接收"],
    ["保温管接收争议审批", "现场负责人", "施工单位提出争议后", "批准则按争议量计算，驳回则恢复确认到货量"],
    ["保温管库管确认", "库管人员", "查看相关记录后", "可批量确认；用于信息知晓和监督，不影响库存"],
    ["保温管使用量填报", "现场负责人", "每日上午8:00前", "按页面显示业务日期填报实际使用量"],
    ["保温管损耗量填报", "现场负责人", "每日上午8:00前", "按页面显示业务日期填报实际损耗量"],
    ["三日计划填报", "现场负责人", "每日上午8:00前", "按页面开放日期维护连续三日滚动计划"],
    ["核心日期维护", "系统管理员", "数据核查后或按自动设置", "核对展示、消耗采集和计划起始日期"],
]
for row, values in zip(quick_table.rows[1:], existing_quick_rows):
    fill_row(row, values)
add_rows(
    quick_table,
    [
        ["管件整车发货", "管厂负责人", "装车完成准备发运时", "同车多种管件逐行填报；同车牌1小时内核对是否合并"],
        ["管件确认到货", "现场负责人", "车辆实际到站并清点后", "逐项填写实际到货数，不得大于发货数"],
        ["管件施工接收", "施工单位人员", "到货确认后", "核对实际到货明细，填写接收说明后确认"],
        ["管件整车归档", "库管人员", "施工接收完成后", "核对车次及明细，执行整车批量归档"],
        ["管件使用量填报", "现场负责人", "页面显示消耗采集日期", "单日提交一次，不得超过现场可用库存"],
        ["基准台账查询", "供给侧、需求侧、管理员", "采购、发运或核对需要时", "按施工标段和物资类型查询，可导出Excel"],
        ["历史查询与导出", "管理人员", "核查或统计需要时", "选择保温管或管件标签，筛选后导出.xlsx台账"],
        ["提交状态核查", "系统管理员", "每日数据核查时", "查看各施工标段最近提交日期、时间和填报人员"],
    ],
)

# 统一用户侧术语和系统名称。排除数字智慧大屏、GIS 等未纳入内容。
replace_text_everywhere(
    doc,
    [
        ("2026年度保温管物流链管理系统", "2026年度保温管、管件物流链管理系统"),
        ("换热站", "施工标段"),
        ("大盘汇总", "综合看板"),
        ("展示/截止日期", "展示日期"),
    ],
)

# 页眉标题位于文本框内，常规段落集合不会遍历到文本框文字；
# 直接更新页眉部件中的文字节点，确保页眉与封面标题一致。
seen_header_parts = set()
for section in doc.sections:
    for story in (section.header, section.first_page_header):
        part = story.part
        if id(part) in seen_header_parts:
            continue
        seen_header_parts.add(id(part))
        for node in part.element.iter(qn("w:t")):
            if node.text:
                node.text = node.text.replace(
                    "2026年度保温管物流链管理系统",
                    "2026年度保温管、管件物流链管理系统",
                )
                node.text = node.text.replace(
                    "年度保温管物流链管理系统",
                    "年度保温管、管件物流链管理系统",
                )
                node.text = node.text.replace(
                    "保温管物流链管理系统",
                    "保温管、管件物流链管理系统",
                )

# 修正原稿常见问题中的漏字。
replace_text_everywhere(doc, [("无法继续使用的温管长度", "无法继续使用的保温管长度")])

# 统一回套原稿中已验证的标题、编号事项样式，消除重新分页后的首字符错位和段落贴合。
for paragraph in doc.paragraphs:
    current = text_of(paragraph)
    if not current:
        continue
    if paragraph.style.name == "Heading 1":
        apply_template_format(paragraph, templates["h1"])
    elif paragraph.style.name == "Heading 2":
        apply_template_format(paragraph, templates["h2"])
    elif paragraph.style.name == "Heading 3":
        apply_template_format(paragraph, templates["h3"])
    elif re.match(r"^（\d+）", current):
        apply_template_format(paragraph, templates["item"])

for paragraph in doc.paragraphs:
    if text_of(paragraph) == "现场负责人确认到货后，系统即按现场负责人确认的实际到货数量计入现场库存。":
        apply_template_format(paragraph, templates["body"])
        break

doc.core_properties.title = "洁净能源集团2026年度保温管、管件物流链管理系统业务操作指南"
doc.core_properties.subject = "V2.0业务操作指南修编稿"
doc.core_properties.comments = "在2026年7月1日V1.0基础上修编；新增管件全流程、历史查询和管理员日常操作。"
update_fields = doc.settings.element.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    doc.settings.element.append(update_fields)
update_fields.set(qn("w:val"), "true")
doc.save(OUTPUT)
print(OUTPUT)
